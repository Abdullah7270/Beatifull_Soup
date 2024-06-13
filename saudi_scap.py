import requests
from bs4 import BeautifulSoup
from docx import Document

url = 'https://en.wikipedia.org/wiki/Economy_of_Saudi_Arabia'

response = requests.get(url)
soup = BeautifulSoup(response.content, 'html.parser')

# Ceate a new document
doc = Document()

# Add title to the document
doc.add_heading('Economy of Saudi Arabia', 0)

# Loop through the paragraphs and add them to the document
texts = soup.find_all('p')
for text in texts:
    doc.add_paragraph(text.text)

# Save the document
doc.save('Economy_of_Saudi_Arabia.docx')

print("The content has been saved to Economy_of_Saudi_Arabia.docx")